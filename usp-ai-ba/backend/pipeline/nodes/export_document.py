"""Node 5 (document mode): write the approved Epic/User Story/Dev Task/Unit Test
Task hierarchy to a .docx file instead of pushing it to Azure DevOps.

Mirrors create_ado_node's content structure and error-handling convention so
the two output modes stay interchangeable via settings.OUTPUT_MODE.
"""
from __future__ import annotations

import logging
import os
import re

from docx import Document
from docx.shared import Pt

from config import settings
from pipeline.state import StoryForgeState

logger = logging.getLogger(__name__)


def _sanitize_filename_part(text: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9]+", "_", text.strip())
    return cleaned.strip("_") or "untitled"


def _next_version(exports_dir: str, base_name: str) -> int:
    """Find the highest existing _V_{n} suffix for base_name and return n + 1."""
    pattern = re.compile(rf"^{re.escape(base_name)}_V_(\d+)\.docx$")
    max_version = 0
    for filename in os.listdir(exports_dir):
        match = pattern.match(filename)
        if match:
            max_version = max(max_version, int(match.group(1)))
    return max_version + 1


def _add_list(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(str(item), style="List Bullet")


def _add_dict_or_list(document: Document, value) -> None:
    if isinstance(value, dict):
        for key, val in value.items():
            document.add_paragraph(f"{key}: {val}", style="List Bullet")
    elif isinstance(value, list):
        _add_list(document, value)
    else:
        document.add_paragraph(str(value))


def _add_story(document: Document, story: dict) -> None:
    document.add_heading(story.get("epic_title", "Untitled Epic"), level=1)

    document.add_heading("User Story", level=2)
    document.add_paragraph(story.get("user_story", ""))

    document.add_heading("Acceptance Criteria", level=2)
    _add_list(document, story.get("acceptance_criteria", []))

    for dev_task in story.get("dev_tasks", []):
        document.add_heading(f"Dev Task: {dev_task.get('title', '')}", level=2)

        document.add_heading("User Story", level=3)
        document.add_paragraph(dev_task.get("user_story", ""))

        document.add_heading("Acceptance Criteria", level=3)
        _add_list(document, dev_task.get("acceptance_criteria", []))

        document.add_heading("Technical Approach", level=3)
        _add_list(document, dev_task.get("technical_approach", []))

        document.add_heading("Affected Components", level=3)
        _add_dict_or_list(document, dev_task.get("affected_components", {}))

        document.add_heading("API Contract", level=3)
        _add_dict_or_list(document, dev_task.get("api_contract", {}))

        document.add_heading("Business Rules", level=3)
        _add_list(document, dev_task.get("business_rules", []))

        document.add_heading("Error Handling", level=3)
        _add_list(document, dev_task.get("error_handling", []))

    for unit_test_task in story.get("unit_test_tasks", []):
        document.add_heading(f"Unit Test Task: {unit_test_task.get('title', '')}", level=2)

        document.add_heading("Test Objective", level=3)
        document.add_paragraph(unit_test_task.get("test_objective", ""))

        document.add_heading("Test Scenarios", level=3)
        for category, scenario_items in unit_test_task.get("test_scenarios", {}).items():
            document.add_paragraph(category, style="List Bullet")
            _add_list(document, scenario_items)

        document.add_heading("Test Data", level=3)
        _add_dict_or_list(document, unit_test_task.get("test_data", {}))

        document.add_heading("Mock Setup", level=3)
        _add_list(document, unit_test_task.get("mock_setup", []))

        document.add_heading("Assertions", level=3)
        _add_list(document, unit_test_task.get("assertions", []))


async def export_document_node(state: StoryForgeState) -> StoryForgeState:
    """Render state["approved_stories"] to a .docx saved under settings.EXPORTS_DIR."""
    job_id = state["job_id"]
    new_errors: list[str] = []
    document_path = ""

    try:
        document = Document()
        style = document.styles["Normal"]
        style.font.size = Pt(10.5)

        document.add_heading(
            f"{state['system_name']} | {state['ppm_number']} | {state['ppm_name']}", level=0
        )

        for story in state["approved_stories"]:
            _add_story(document, story)

        os.makedirs(settings.EXPORTS_DIR, exist_ok=True)
        base_name = "_".join(
            _sanitize_filename_part(part)
            for part in (state["ppm_number"], state["ppm_name"], state["system_name"])
        )
        version = _next_version(settings.EXPORTS_DIR, base_name)
        document_path = os.path.join(settings.EXPORTS_DIR, f"{base_name}_V_{version}.docx")
        document.save(document_path)
    except Exception as exc:  # noqa: BLE001 - record and surface, don't crash the graph
        logger.exception("Failed to export document for job %s", job_id)
        new_errors.append(f"export_document_node: {exc}")
        document_path = ""

    return {
        **state,
        "document_path": document_path,
        "errors": state["errors"] + new_errors,
        "status": "done" if not new_errors else "error",
    }
