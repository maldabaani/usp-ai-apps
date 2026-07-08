"""One-time ingestion of user-facing documents -- PDF and Word (.docx) only --
into the ``sf_user_manuals`` ChromaDB collection.

Deliberately narrow: earlier this also accepted Markdown and Confluence-HTML
exports, but that let non-document files (release notes, README-style
Markdown, etc.) leak into the manuals collection alongside real business
documents. Restricted back to just PDF/Word so "Ingest Documents" only ever
indexes actual documents -- source code and code-adjacent files (.js/.json/
.md/etc.) are handled exclusively by the separate "Ingest Code" flow
(ingest_code.py) and were never accepted here to begin with.
"""
from __future__ import annotations

import hashlib
import logging
import time
from pathlib import Path

import docx
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from pypdf import PdfReader

from config import settings
from ingestion.chroma_client import delete_by_source, get_vector_store, list_distinct_sources

logger = logging.getLogger(__name__)

CHUNK_SIZE_TOKENS = 1500
CHUNK_OVERLAP_TOKENS = 150
# Approximate characters-per-token ratio used to translate the token-based
# limits in the spec into character-based limits for the text splitter.
# Tightened from 4 -- see ingest_code.py's CHARS_PER_TOKEN comment for the
# full story: Ollama's nomic-embed-text has a real, unconfigurable 2048-token
# embedding ceiling (num_ctx overrides aren't honored for this model), and a
# live ingestion run showed "4x-sized" chunks tokenizing up to that boundary.
# Prose chunks here are less likely to hit code's worst-case ratio, but using
# the same conservative value keeps both ingestion paths under one shared,
# easy-to-reason-about ceiling instead of two different unverified guesses.
CHARS_PER_TOKEN = 3

_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE_TOKENS * CHARS_PER_TOKEN,
    chunk_overlap=CHUNK_OVERLAP_TOKENS * CHARS_PER_TOKEN,
    separators=["\n\n", "\n", ". ", " ", ""],
)

# Suffix -> ("format" metadata tag, extractor). Checked via a dict keyed on
# path.suffix.lower() rather than a chain of if/elif branches.
_FORMAT_BY_SUFFIX = {
    ".pdf": "pdf",
    ".docx": "docx",
}
SUPPORTED_EXTENSIONS = frozenset(_FORMAT_BY_SUFFIX)


def _extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    pages_text = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(f"[Page {page_number}]\n{text}")
    return "\n\n".join(pages_text)


def _extract_docx_text(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_text(path: Path) -> str:
    format_tag = _FORMAT_BY_SUFFIX[path.suffix.lower()]
    if format_tag == "pdf":
        return _extract_pdf_text(path)
    return _extract_docx_text(path)


def _document_id(relative_source: str, chunk_index: int) -> str:
    """Deterministic per-chunk ID, matching ingest_code.py's fix for the same
    duplicate-on-rerun bug this module's own docstring used to document as a
    known limitation."""
    key = f"{relative_source}::{chunk_index}"
    return hashlib.sha256(key.encode("utf-8")).hexdigest()[:40]


def _chunk_document(path: Path, folder_path: Path) -> list[Document]:
    full_text = _extract_text(path)
    if not full_text.strip():
        logger.warning("No extractable text found in %s", path)
        return []

    relative_source = str(path.relative_to(folder_path))
    format_tag = _FORMAT_BY_SUFFIX[path.suffix.lower()]
    chunks = _splitter.split_text(full_text)
    ingested_at = time.time()
    return [
        Document(
            page_content=chunk,
            metadata={
                "source": relative_source,
                "type": "user_manual",
                "module": "N/A",
                "layer": "documentation",
                "class_name": "N/A",
                "language": "N/A",
                "chunk_index": index,
                "format": format_tag,
                "ingested_at": ingested_at,
            },
        )
        for index, chunk in enumerate(chunks)
    ]


async def ingest_documents(
    folder_path: str,
    progress_callback=None,
    *,
    enable_llm_summary: bool | None = None,
    max_concurrency: int | None = None,
) -> dict:
    """Chunk and embed every supported document found (recursively) under
    ``folder_path`` -- PDF and Word (.docx) only.

    Stores all chunks into the ``sf_user_manuals`` collection. Safe to re-run:
    deterministic per-chunk IDs make an unchanged file's re-add a no-op
    upsert, each file's prior chunk set is cleared before its fresh one is
    added (so a changed file's stale chunks don't linger), and a file removed
    from the folder since the last run has its chunks purged too (diffed
    against what the collection already has, since nothing else here
    persists prior runs).

    After tier 1 (mechanical chunking) finishes, optionally runs tier 2 (see
    ingestion/enrichment/enrich_documents.py): a per-document LLM-synthesized
    business-rule summary, embedded alongside the raw chunks above.
    ``enable_llm_summary`` defaults to settings.INGEST_LLM_SUMMARY_ENABLED
    when not given explicitly (a per-request override, matching
    ingest_code.py's own precedent).
    """
    # Imported lazily (not at module top) to avoid a circular import --
    # enrich_documents.py itself imports _extract_text/_splitter/
    # _FORMAT_BY_SUFFIX from this module at its own top level.
    from ingestion.enrichment import enrich_documents as enrich_documents_module

    folder = Path(folder_path)
    if not folder.is_dir():
        raise FileNotFoundError(f"Documents folder not found: {folder_path}")

    doc_paths = sorted(
        path for path in folder.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )
    vector_store = get_vector_store("manuals")

    current_sources = {str(path.relative_to(folder)) for path in doc_paths}
    previously_seen = await list_distinct_sources("manuals")
    for removed_source in previously_seen - current_sources:
        await delete_by_source("manuals", removed_source)

    total_files = len(doc_paths)
    total_chunks = 0
    errors: list[str] = []
    file_records: list[dict] = []

    for index, doc_path in enumerate(doc_paths, start=1):
        relative_source = str(doc_path.relative_to(folder))
        try:
            documents = _chunk_document(doc_path, folder)
            if documents:
                await delete_by_source("manuals", relative_source)
                ids = [_document_id(relative_source, doc.metadata["chunk_index"]) for doc in documents]
                await vector_store.aadd_documents(documents, ids=ids)
                total_chunks += len(documents)
                file_records.append({"path": relative_source, "status": "success", "chunks": len(documents)})
            else:
                file_records.append({"path": relative_source, "status": "skipped", "reason": "no_content_extracted"})
        except Exception as exc:  # noqa: BLE001 - surfaced to caller via errors list
            logger.exception("Failed to ingest %s", doc_path)
            errors.append(f"{doc_path}: {exc}")
            file_records.append({"path": relative_source, "status": "error", "reason": str(exc)})

        if progress_callback:
            await progress_callback(
                index, total_files, phase="chunking", partial_result={"files": file_records.copy()}
            )

    resolved_enable_llm_summary = (
        settings.INGEST_LLM_SUMMARY_ENABLED if enable_llm_summary is None else enable_llm_summary
    )
    enrichment_result = await enrich_documents_module.enrich_documents(
        folder,
        doc_paths,
        enabled=resolved_enable_llm_summary,
        max_concurrency=max_concurrency or enrich_documents_module.DEFAULT_MAX_CONCURRENCY,
        progress_callback=progress_callback,
    )

    return {
        "files_processed": total_files,
        "chunks_indexed": total_chunks,
        "errors": errors + enrichment_result["errors"],
        "llm_summary_enabled": enrichment_result["enabled"],
        "files_summarized": enrichment_result["files_summarized"],
        "files_skipped_unchanged": enrichment_result["files_skipped_unchanged"],
        "files": file_records,
        "enrichment_files": enrichment_result["files"],
    }
