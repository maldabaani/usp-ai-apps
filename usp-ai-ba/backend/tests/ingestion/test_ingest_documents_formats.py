"""Covers ingest_documents.py's supported-format allow-list: PDF and Word
(.docx) only. Markdown and Confluence-HTML-export support (added in an
earlier phase) was deliberately removed -- it let non-document files leak
into the manuals collection -- and .js/.json were never accepted here at all
(those belong exclusively to the separate "Ingest Code" flow). This confirms
that narrowing, not just the extraction logic for the two formats that remain.

Uses a fake vector store standing in for `langchain_chroma.Chroma`, matching
this codebase's mocked-client testing convention (see tests/ingestion/test_dedup.py).
"""
from __future__ import annotations

import asyncio
import uuid

import docx
import pytest

from ingestion import chroma_client, ingest_documents


class _FakeVectorStore:
    def __init__(self):
        self.docs: dict[str, object] = {}

    async def aadd_documents(self, documents, ids=None):
        ids = ids or [str(uuid.uuid4()) for _ in documents]
        for id_, doc in zip(ids, documents):
            self.docs[id_] = doc
        return ids

    async def adelete(self, ids=None, where=None):
        if where:
            for id_ in [i for i, doc in self.docs.items() if doc.metadata.get("source") == where.get("source")]:
                del self.docs[id_]
        elif ids:
            for id_ in ids:
                self.docs.pop(id_, None)

    def get(self, include=None):
        return {"metadatas": [doc.metadata for doc in self.docs.values()]}


@pytest.fixture
def fake_manuals_store(monkeypatch):
    store = _FakeVectorStore()

    def fake_get_vector_store(collection_key: str):
        return store

    monkeypatch.setattr(chroma_client, "get_vector_store", fake_get_vector_store)
    monkeypatch.setattr(ingest_documents, "get_vector_store", fake_get_vector_store)
    return store


def _write_docx(path, paragraphs: list[str]) -> None:
    document = docx.Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(path))


def test_extract_docx_text_includes_paragraphs_and_table_cells(tmp_path):
    path = tmp_path / "doc.docx"
    document = docx.Document()
    document.add_paragraph("Some paragraph text.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Key"
    table.rows[0].cells[1].text = "Value"
    document.save(str(path))

    text = ingest_documents._extract_docx_text(path)

    assert "Some paragraph text." in text
    assert "Key | Value" in text


@pytest.mark.parametrize("suffix", [".md", ".markdown", ".html", ".htm", ".doc", ".js", ".json", ".txt"])
def test_unsupported_extensions_are_excluded(tmp_path, fake_manuals_store, suffix):
    (tmp_path / f"file{suffix}").write_text("some content that would otherwise be ingested")
    _write_docx(tmp_path / "spec.docx", ["Word document content."])

    result = asyncio.run(ingest_documents.ingest_documents(str(tmp_path)))

    # Only the .docx file is picked up -- the other extension is invisible to
    # the walk entirely, not skipped-with-a-reason.
    assert result["files_processed"] == 1
    assert [f["path"] for f in result["files"]] == ["spec.docx"]


def test_mixed_pdf_and_docx_folder_tags_correct_format_per_source(tmp_path, fake_manuals_store, monkeypatch):
    _write_docx(tmp_path / "spec.docx", ["Word document content."])
    (tmp_path / "manual.pdf").write_text("not a real pdf, but extraction is stubbed below")

    monkeypatch.setattr(ingest_documents, "_extract_pdf_text", lambda path: "Extracted PDF body text.")

    result = asyncio.run(ingest_documents.ingest_documents(str(tmp_path)))

    assert result["files_processed"] == 2
    assert not result["errors"]

    formats_by_source = {
        doc.metadata["source"]: doc.metadata["format"] for doc in fake_manuals_store.docs.values()
    }
    assert formats_by_source == {"manual.pdf": "pdf", "spec.docx": "docx"}
